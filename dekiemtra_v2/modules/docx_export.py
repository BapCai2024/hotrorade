# -*- coding: utf-8 -*-
from __future__ import annotations

import io
import re
from typing import Any, Dict, List, Tuple

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _set_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)


def _split_answer(content: str) -> Tuple[str, str]:
    content = (content or "").strip()
    m = re.search(r"(?im)^\s*đáp\s*án\s*:\s*(.+)$", content)
    if not m:
        return content, ""
    ans = m.group(1).strip()
    stem = re.sub(r"(?im)^\s*đáp\s*án\s*:\s*.+$", "", content).strip()
    return stem, ans


def create_exam_docx(
    school_name: str,
    subject: str,
    grade: str,
    exam_term: str,
    exam_list: List[Dict[str, Any]],
    include_answers: bool,
) -> io.BytesIO:
    doc = Document()
    _set_font(doc)

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(7)
    table.columns[1].width = Cm(9)

    c1 = table.cell(0, 0).paragraphs[0]
    r1 = c1.add_run((school_name or "").upper())
    r1.bold = True
    c1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    c2 = table.cell(0, 1).paragraphs[0]
    r2 = c2.add_run(f"{(exam_term or 'ĐỀ KIỂM TRA').upper()}\nMÔN: {subject.upper()} — {grade.upper()}")
    r2.bold = True
    c2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    title = doc.add_heading("ĐỀ BÀI", level=1)
    if title.runs:
        title.runs[0].font.name = "Times New Roman"

    answers: List[str] = []
    for idx, q in enumerate(exam_list, start=1):
        stem, ans = _split_answer(q.get("content", ""))

        already_numbered = bool(re.match(r"(?is)^\s*câu\s+\d+", stem.strip()))
        if not already_numbered:
            p = doc.add_paragraph()
            run = p.add_run(f"Câu {idx} ({q.get('points','')} điểm): ")
            run.bold = True

        for line in stem.splitlines():
            if line.strip():
                doc.add_paragraph(line.strip())
        doc.add_paragraph()

        if include_answers and ans:
            answers.append(f"Câu {idx}: {ans}")

    if include_answers:
        doc.add_page_break()
        h2 = doc.add_heading("ĐÁP ÁN", level=1)
        if h2.runs:
            h2.runs[0].font.name = "Times New Roman"
        for a in answers:
            doc.add_paragraph(a)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def create_matrix_docx(subject: str, grade: str, exam_list: List[Dict[str, Any]]) -> io.BytesIO:
    doc = Document()
    _set_font(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"BẢNG MA TRẬN (BẢN ĐẶC TẢ) — {subject.upper()} {grade.upper()}")
    r.bold = True

    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    hdrs = ["STT", "Chủ đề", "Bài học", "YCCĐ", "Dạng", "Mức", "Điểm"]
    for i, h in enumerate(hdrs):
        table.rows[0].cells[i].text = h

    for idx, q in enumerate(exam_list, start=1):
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = str(q.get("topic", ""))
        row[2].text = str(q.get("lesson", ""))
        row[3].text = str(q.get("yccd", ""))
        row[4].text = str(q.get("type", ""))
        row[5].text = str(q.get("level", ""))
        row[6].text = str(q.get("points", ""))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
