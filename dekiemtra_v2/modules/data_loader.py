# -*- coding: utf-8 -*-
from __future__ import annotations

import io
import re
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd
from docx import Document

try:
    import pypdf  # type: ignore
    PDF_ENABLED = True
except Exception:
    PDF_ENABLED = False

MAX_FILE_TEXT_CHARS = 60_000
MAX_XLSX_ROWS_FOR_PROMPT = 200


def _truncate(s: str, max_chars: int) -> str:
    s = s or ""
    return s if len(s) <= max_chars else s[:max_chars] + "\n\n[...ĐÃ CẮT BỚT...]"


def extract_text_from_upload(filename: str, data: bytes) -> Tuple[Optional[str], Optional[str]]:
    """Đọc file ma trận (xlsx/docx/pdf) => text để đưa vào prompt (Tab 1)."""
    try:
        name = (filename or "").lower()
        bio = io.BytesIO(data)

        if name.endswith(".xlsx"):
            df = pd.read_excel(bio)
            df2 = df.head(MAX_XLSX_ROWS_FOR_PROMPT).copy()
            df2 = df2.fillna("").astype(str)
            text = df2.to_csv(index=False)
            return _truncate(text, MAX_FILE_TEXT_CHARS), None

        if name.endswith(".docx"):
            doc = Document(bio)
            parts: List[str] = []
            for table in doc.tables:
                for row in table.rows:
                    parts.append(" | ".join((c.text or "").strip() for c in row.cells))
            if not parts:
                parts = [p.text for p in doc.paragraphs if (p.text or "").strip()]
            text = "\n".join(parts)
            return _truncate(text, MAX_FILE_TEXT_CHARS), None

        if name.endswith(".pdf"):
            if not PDF_ENABLED:
                return None, "Thiếu thư viện pypdf. Cài: pip install pypdf"
            reader = pypdf.PdfReader(bio)
            pages: List[str] = []
            for page in reader.pages:
                t = page.extract_text() or ""
                if t.strip():
                    pages.append(t)
            text = "\n".join(pages).strip()
            if not text:
                return None, "Không trích xuất được text PDF (có thể là file scan ảnh)."
            return _truncate(text, MAX_FILE_TEXT_CHARS), None

        return None, "Định dạng không hỗ trợ (chỉ xlsx/docx/pdf)."
    except Exception as e:
        return None, f"Lỗi đọc file: {e}"


def _normalize_header(s: str) -> str:
    s = (s or "").strip().lower()
    s = re.sub(r"\s+", " ", s)
    return s


def build_nested_curriculum(df: pd.DataFrame) -> Dict[str, Any]:
    nested: Dict[str, Any] = {}
    if df is None or df.empty:
        return nested

    for _, r in df.iterrows():
        lop = (r.get("lop") or "").strip() or "Khác"
        mon = (r.get("mon") or "").strip() or "Khác"
        hk = (r.get("hoc_ky") or "").strip() or "Khác"
        cd = (r.get("chu_de") or "").strip() or "Khác"
        bai = (r.get("bai") or "").strip() or ""

        nested.setdefault(lop, {}).setdefault(mon, {}).setdefault(hk, {}).setdefault(cd, [])
        if bai and bai not in nested[lop][mon][hk][cd]:
            nested[lop][mon][hk][cd].append(bai)
    return nested


def load_curriculum_from_docx(docx_bytes: bytes) -> Tuple[pd.DataFrame, Dict[str, Any], str]:
    """
    Đọc DOCX dạng bảng (Học kì/Lớp/Môn/Chủ đề/Bài...) và chuẩn hoá:
    - df: bảng phẳng
    - nested: dict để dropdown
    - warn: cảnh báo thiếu cột (Tiết, YCCĐ, Bộ sách...)
    """
    doc = Document(io.BytesIO(docx_bytes))

    rows: List[List[str]] = []
    for table in doc.tables:
        for row in table.rows:
            rows.append([(cell.text or "").strip() for cell in row.cells])

    if not rows:
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        df = pd.DataFrame({"raw": lines})
        return df, {}, "DOCX không có bảng. Khuyến nghị dùng file có bảng hoặc dùng Excel làm nguồn chuẩn."

    header_idx = None
    for i, r in enumerate(rows[:30]):
        joined = " ".join(_normalize_header(x) for x in r)
        if "lớp" in joined and "môn" in joined:
            header_idx = i
            break
    if header_idx is None:
        header_idx = 0

    header = rows[header_idx]
    cols = [_normalize_header(c) for c in header]

    data_rows: List[List[str]] = []
    for r in rows[header_idx + 1 :]:
        if len(r) < 3:
            continue
        jr = " ".join(_normalize_header(x) for x in r)
        if "lớp" in jr and "môn" in jr and "chủ đề" in jr:
            continue
        data_rows.append(r + [""] * (len(cols) - len(r)))

    df = pd.DataFrame(data_rows, columns=[c if c else f"col_{i}" for i, c in enumerate(cols)])

    col_map = {}
    for c in df.columns:
        if "học kì" in c or "hoc ki" in c:
            col_map[c] = "hoc_ky"
        elif c == "lớp" or "lop" in c:
            col_map[c] = "lop"
        elif "môn" in c or "mon" in c:
            col_map[c] = "mon"
        elif "chủ đề" in c or "chu de" in c:
            col_map[c] = "chu_de"
        elif "tên bài" in c or "bài học" in c or "bai hoc" in c:
            col_map[c] = "bai"
        elif "tiết" in c or "tiet" in c:
            col_map[c] = "tiet"
        elif "yccđ" in c or "yccd" in c:
            col_map[c] = "yccd"
        elif "bộ sách" in c or "bo sach" in c:
            col_map[c] = "bo_sach"

    df = df.rename(columns=col_map)

    for must in ["hoc_ky", "lop", "mon", "chu_de", "bai"]:
        if must not in df.columns:
            df[must] = ""

    for c in df.columns:
        df[c] = df[c].astype(str).str.strip()

    missing = [c for c in ["bo_sach", "tiet", "yccd"] if c not in df.columns]
    warn = ""
    if missing:
        warn = "Thiếu cột: " + ", ".join(missing) + ". Bạn vẫn dùng dropdown Chủ đề/Bài, nhưng để chuẩn CT2018 nên bổ sung (khuyến nghị Excel)."

    nested = build_nested_curriculum(df)
    return df, nested, warn


def load_sample_curriculum() -> Tuple[pd.DataFrame, Dict[str, Any]]:
    sample = [
        {"hoc_ky": "Học kì I", "lop": "Lớp 5", "mon": "Khoa học", "chu_de": "Chất và sự biến đổi", "bai": "Hỗn hợp và dung dịch", "tiet": "1", "yccd": ""},
        {"hoc_ky": "Học kì I", "lop": "Lớp 5", "mon": "Khoa học", "chu_de": "Chất và sự biến đổi", "bai": "Tách các chất trong hỗn hợp", "tiet": "1", "yccd": ""},
        {"hoc_ky": "Học kì I", "lop": "Lớp 4", "mon": "Lịch sử và Địa lí", "chu_de": "Địa phương em", "bai": "Thiên nhiên địa phương", "tiet": "1", "yccd": ""},
    ]
    df = pd.DataFrame(sample)
    return df, build_nested_curriculum(df)
